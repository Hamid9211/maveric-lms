"""
Extract plain text from an uploaded article document (txt / md / pdf / docx).

The extracted text is treated exactly like text a teacher types into the
article box — it is returned to the caller and ultimately persisted in the
Chapter.article_content column (Supabase/Postgres). No raw file is stored; the
text itself is the source of truth used for embeddings and quiz generation.
"""
import io
import os
from fastapi import HTTPException

ALLOWED_DOC_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}
MAX_DOC_SIZE = 15 * 1024 * 1024  # 15 MB


def extract_text_from_document(filename: str, data: bytes) -> str:
    """Return the plain text contained in an uploaded document.

    Raises HTTPException(400) for unsupported types, oversized files, unreadable
    documents, or documents with no extractable text (e.g. scanned PDFs).
    """
    ext = os.path.splitext(filename or "")[1].lower()
    if ext not in ALLOWED_DOC_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext or 'unknown'}'. Allowed: .txt, .md, .pdf, .docx",
        )
    if len(data) > MAX_DOC_SIZE:
        raise HTTPException(status_code=400, detail="File is too large. Maximum size is 15MB.")

    try:
        if ext in (".txt", ".md"):
            text = data.decode("utf-8", errors="replace")
        elif ext == ".pdf":
            text = _extract_pdf(data)
        else:  # .docx
            text = _extract_docx(data)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Could not read the document: {exc}")

    text = text.strip()
    if not text:
        raise HTTPException(
            status_code=400,
            detail="No readable text found in the document. If it is a scanned PDF, paste the text manually.",
        )
    return text


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = [(page.extract_text() or "") for page in reader.pages]
    return "\n\n".join(parts)


def _extract_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    # Include table cell text so tabular notes aren't lost.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)
